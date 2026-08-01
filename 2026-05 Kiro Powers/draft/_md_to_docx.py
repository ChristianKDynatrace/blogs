"""One-off markdown -> docx converter for the Kiro Powers blog V3."""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


INLINE_RE = re.compile(
    r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))'
)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, text, base_italic=False):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.italic = base_italic
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.italic = base_italic
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        elif token.startswith("["):
            mlink = re.match(r'\[([^\]]+)\]\(([^)]+)\)', token)
            if mlink:
                add_hyperlink(paragraph, mlink.group(2), mlink.group(1))
            else:
                run = paragraph.add_run(token)
                run.italic = base_italic
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.italic = base_italic


def convert(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        # Markdown table: pipe-delimited rows with separator line
        if (stripped.startswith("|") and i + 1 < len(lines)
                and re.match(r'^\s*\|[\s:|\-]+\|\s*$', lines[i+1])):
            def parse_cells(ln):
                s = ln.strip()
                if s.startswith("|"):
                    s = s[1:]
                if s.endswith("|"):
                    s = s[:-1]
                return [c.strip() for c in s.split("|")]
            header = parse_cells(lines[i])
            i += 2  # skip separator
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_cells(lines[i]))
                i += 1
            tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
            tbl.style = "Light Grid Accent 1"
            for c, txt in enumerate(header):
                cell_p = tbl.cell(0, c).paragraphs[0]
                add_inline_runs(cell_p, txt)
                for r in cell_p.runs:
                    r.bold = True
            for r_idx, row in enumerate(rows):
                for c, txt in enumerate(row):
                    cell_p = tbl.cell(r_idx + 1, c).paragraphs[0]
                    add_inline_runs(cell_p, txt)
            continue

        # Image
        m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if m:
            alt, src = m.group(1), m.group(2)
            img_path = (md_path.parent / src).resolve()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_path.exists():
                p.add_run().add_picture(str(img_path), width=Inches(5.5))
            else:
                p.add_run(f"[Image not found: {src}]").italic = True
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                content = lines[i].lstrip()[1:].lstrip()
                quote_lines.append(content)
                i += 1
            # Group into paragraphs separated by empty quote-lines
            paragraphs = []
            current = []
            for q in quote_lines:
                if q == "":
                    if current:
                        paragraphs.append(" ".join(current))
                        current = []
                else:
                    current.append(q)
            if current:
                paragraphs.append(" ".join(current))
            for para in paragraphs:
                p = doc.add_paragraph(style="Intense Quote") if "Intense Quote" in [s.name for s in doc.styles] else doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                add_inline_runs(p, para)
            continue

        # Numbered list
        m = re.match(r'^\d+\.\s+(.*)$', stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m.group(1))
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- "):
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, content)
            # collect continuation indented lines (sub-content under the bullet)
            i += 1
            while i < len(lines):
                nxt = lines[i]
                if nxt.startswith("  ") and nxt.strip():
                    # treat as separate indented paragraph under bullet
                    sub = nxt.strip()
                    sub_m = re.match(r'`([^`]+)`', sub)
                    sp = doc.add_paragraph()
                    sp.paragraph_format.left_indent = Inches(0.5)
                    add_inline_runs(sp, sub)
                    i += 1
                elif nxt.strip() == "":
                    # peek further
                    if i + 1 < len(lines) and lines[i+1].startswith("  "):
                        i += 1
                    else:
                        break
                else:
                    break
            continue

        # Regular paragraph (gather consecutive lines, preserve as soft line breaks)
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,3} |!\[|>|\d+\.\s|-\s)', lines[i].lstrip()):
            para_lines.append(lines[i].strip())
            i += 1
        joined = "\n".join(para_lines)
        # detect fully italicized caption (single line)
        if len(para_lines) == 1 and joined.startswith("*") and joined.endswith("*") and not joined.startswith("**"):
            inner = joined[1:-1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, inner, base_italic=True)
        else:
            p = doc.add_paragraph()
            for idx, ln in enumerate(para_lines):
                if idx > 0:
                    br = OxmlElement("w:br")
                    p._p[-1].append(br) if False else None
                    run = p.add_run()
                    run._r.append(OxmlElement("w:br"))
                add_inline_runs(p, ln)

    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    md = Path(sys.argv[1])
    out = Path(sys.argv[2])
    convert(md, out)
